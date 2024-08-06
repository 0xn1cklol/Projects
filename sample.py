'''
Programmer: Nick Ricci
Date: 8/6/2024
Time Started: 11:03am
Time Finished: 12:51pm

Purpose: 
    - Create a program that imports excel file and gets table contents
    - Take those contents and move them to a word document
    - Give output that the write was performed
    - Write a test function to verify the write

Verdict:
    - Script successfuly imports excel file and creates a new docx file
    - Script successfully reads excel file and transfers data to docx file
    - Script successfully outputs informative information to user
    - Script successfully verfies transfer of file contents 

Conclusion:
      Had to do a bit of researching around to get things to work
      Heres where I struggled a bit:
        - I couldnt open a docx file that already existed, so I had to create one in the script
        - Wasnt sure how you wanted the format into the docx to be so I just left it the same
        - FUTURE IMPLEMENTATION would obviously be to allow users to choose their own files, name their own files, format the files etc.. this was kept simple for now until further instructions

'''

import sys                                      # used to solely exit script 
from openpyxl import load_workbook              # used to open excel file to get table from 
from docx import Document                       # used to open a word document that we will eventually write to 


docx_file = Document()                          # opening a blank word document

excel_file = "sample.xlsx"                      # name of excel file
workbook = load_workbook(filename=excel_file)   # opening the excel file 
sheet = workbook.active                         # defining the workbook

WrittenData = []                                # a list of the data that was moved from excel to word doc

for row in sheet.iter_rows(values_only=True):   # for loop to get all contents from table 
    a = 0                                       # initialize variable
    length = len(row)                           # getting the length of the row
    line = ""                                   # another initalize variable but for strings
    for item in row:                            # starting a for loop that gets each single item inside a row
        line = line + "     " + item            # creating a single line with each of the items from the row
        if a == length:                         # if we have all items from the row
            a = 0                               # reset the initialize variable
            break                               # break the second for loop
        else:                                   # else statement
            continue                            # continue if conditions not met
    
    paragraph = docx_file.add_paragraph()       # initialize the paragraph feature
    paragraph.add_run(line)                     # Adding the line to docx file
    WrittenData.append(line)                    # Adding data to a list so we can verify the transfer later
#

def VerifyWrite(filename,list):                         # Creating a function to verify content transfer
    verify_file = Document(filename)                    # Opening the newly made docx file
    for paragraph in verify_file.paragraphs:            # basically saying for each line in the file
        if paragraph.text in list:                      # checks if that document line is in the verified list 
            continue                                    # continue statement if condition is met
        else:                                           # else
            print ("Script didnt transfer correctly")   # something went wrong
            sys.exit()                                  # exits the script from continuing 
    print ("Transfer Looked Good :)")                 # informing user the script transfer looked good


docx_file.save("File.docx")                                                 # saving the docx file
print ("File has been saved to 'File.docx' in your current directory!")     # output to let the user know where the file has been saved
verify = input("Would you like to verify the write?[Y/n] > ")               # verify the write function question

if verify == "Y" or verify == "y":                      # checking to see if they answered with Y or y. This probs could have been done simpler using .upper() or .lower()
    VerifyWrite("File.docx",WrittenData)                # starting the VerifyWrite function if user selected 'y' 



